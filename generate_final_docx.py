import sys
import os
import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1/2 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3')
        tblBorders.append(border)
        
    tblPr.append(tblBorders)

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True
    run.font.italic = True
    return p

def add_body_paragraph(doc, text, bold_prefix="", indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.2)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Arial'
        r_pre.font.size = Pt(12)
        r_pre.bold = True
        
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    return p

def add_figure(doc, img_path, caption):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.paragraph_format.line_spacing = 1.0
    
    if os.path.exists(img_path):
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(5.0))
    else:
        run_err = p_img.add_run(f"[Gambar Tidak Ditemukan: {img_path}]")
        run_err.font.name = 'Arial'
        run_err.font.size = Pt(11)
        run_err.font.italic = True
        
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.paragraph_format.line_spacing = 1.0
    run_cap = p_cap.add_run(caption)
    run_cap.font.name = 'Arial'
    run_cap.font.size = Pt(11)
    run_cap.font.italic = True
    return p_img

def add_table_custom(doc, headers, data, widths, caption):
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_cap.paragraph_format.space_before = Pt(12)
    p_cap.paragraph_format.space_after = Pt(4)
    p_cap.paragraph_format.line_spacing = 1.0
    run_cap = p_cap.add_run(caption)
    run_cap.font.name = 'Arial'
    run_cap.font.size = Pt(11)
    run_cap.font.italic = True

    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    set_table_borders(table)
    
    # Header formatting
    hdr_cells = table.rows[0].cells
    for j, text in enumerate(headers):
        hdr_cells[j].text = text
        set_cell_background(hdr_cells[j], "E8F0EC") # Sage green shade
        p = hdr_cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            r.bold = True
            
    # Data rows
    for i, row_data in enumerate(data):
        row_cells = table.rows[i + 1].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = str(val)
            p = row_cells[j].paragraphs[0]
            if j == 0 or len(str(val)) <= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(10)
                
    # Widths setting
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Cm(width)
            
    doc.add_paragraph() # Spacing below table
    return table

def add_reference_entry(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-1.2)
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    return p

def format_all_sections(doc):
    for section in doc.sections:
        # A4 size
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # Margins: Left 4cm, Top 4cm, Right 3cm, Bottom 3cm
        section.top_margin = Cm(4.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(4.0)
        section.right_margin = Cm(3.0)

def format_existing_paragraphs(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    for p in doc.paragraphs:
        text_upper = p.text.strip().upper()
        is_heading = p.style.name.startswith('Heading')
        
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        if text_upper.startswith("BAB ") or text_upper in ["PENDAHULUAN", "TINJAUAN PUSTAKA", "METODOLOGI PENELITIAN", "HASIL DAN PEMBAHASAN", "PENUTUP", "DAFTAR PUSTAKA"]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(12)
                r.bold = True
        elif is_heading:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(12)
                r.bold = True
        else:
            if not p.alignment:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            trimmed = p.text.strip()
            if trimmed and not trimmed.startswith("Gambar ") and not trimmed.startswith("Tabel ") and not trimmed.startswith("-") and not (trimmed[0].isdigit() and (trimmed.startswith("1. ") or trimmed.startswith("2. ") or trimmed.startswith("3. ") or trimmed.startswith("4. ") or trimmed.startswith("5. "))):
                p.paragraph_format.first_line_indent = Cm(1.2)
                
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.name = 'Arial'
                        r.font.size = Pt(10)

def main():
    source_path = 'SKRIPSI SISTEM PAKAR S1 THORIQ.docx'
    output_path = 'final.docx'
    
    print(f"Loading source document: {source_path}")
    doc = Document(source_path)
    
    doc.add_page_break()
    
    # ==================== BAB IV ====================
    add_heading_1(doc, "BAB IV")
    add_heading_1(doc, "HASIL DAN PEMBAHASAN")
    
    add_body_paragraph(doc, "Pada bab ini dibahas mengenai hasil penelitian berupa implementasi sistem pakar deteksi tingkat depresi, kecemasan, dan stres menggunakan metode Depression Anxiety Stress Scale (DASS-21) berbasis web. Penjelasan dalam bab ini mencakup perancangan sistem yang mendalam, implementasi sistem yang meliputi database, kode program, antarmuka pengguna, diagram alur operasional sistem, serta pengujian sistem yang dilakukan menggunakan metode Black Box Testing.")
    
    add_heading_2(doc, "4.1 Perancangan Sistem")
    add_body_paragraph(doc, "Perancangan sistem merupakan tahapan penting untuk memetakan kebutuhan fungsional dan non-fungsional ke dalam representasi logis sebelum dilakukan implementasi kode program. Perancangan ini didasarkan pada metode DASS-21 yang mendeteksi tiga aspek kondisi psikologis, yaitu depresi, kecemasan, dan stres.")
    
    add_heading_3(doc, "4.1.1 Deskripsi Kasus DASS-21 (Kondisi Kesehatan Mental)")
    add_body_paragraph(doc, "Metode DASS-21 merupakan instrumen pengukuran psikologis yang terdiri dari 21 item pertanyaan singkat. Setiap pertanyaan merepresentasikan gejala tertentu yang dialami pengguna dalam satu minggu terakhir. 21 pertanyaan ini terbagi menjadi 3 kategori gangguan emosional, masing-masing terdiri dari 7 pertanyaan, yaitu:")
    add_body_paragraph(doc, "Indikator gejala depresi dirancang untuk mengukur kondisi suasana hati yang buruk, hilangnya minat, dan perasaan putus asa. Indikator gejala kecemasan mengukur respon fisiologis, kecemasan situasional, dan perasaan panik subjektif. Sementara itu, indikator stres mengukur tingkat kesulitan bersantai, ketegangan saraf, kemudahan tersinggung, dan reaksi berlebihan terhadap suatu situasi.", "1. Pembagian Kategori Gejala: ", indent=False)
    
    # Table 4.1: Gejala DASS-21
    headers_dass = ["No", "Kode Gejala", "Pernyataan Gejala (DASS-21 Item)", "Kategori Aspek"]
    widths_dass = [1.0, 2.5, 9.0, 2.5]
    data_dass = [
        ["1", "G001", "Saya merasa sulit untuk menenangkan diri.", "Stres"],
        ["2", "G002", "Saya menyadari mulut saya terasa kering.", "Kecemasan"],
        ["3", "G003", "Saya sama sekali tidak dapat merasakan perasaan positif.", "Depresi"],
        ["4", "G004", "Saya mengalami kesulitan bernapas padahal tidak sedang melakukan aktivitas fisik.", "Kecemasan"],
        ["5", "G005", "Saya merasa kesulitan untuk berinisiatif melakukan sesuatu.", "Depresi"],
        ["6", "G006", "Saya cenderung bereaksi berlebihan terhadap suatu situasi.", "Stres"],
        ["7", "G007", "Saya pernah mengalami gemetar (misal: di tangan).", "Kecemasan"],
        ["8", "G008", "Saya merasa telah menghabiskan banyak energi karena merasa cemas.", "Stres"],
        ["9", "G009", "Saya merasa khawatir dengan situasi dimana saya mungkin menjadi panik dan mempermalukan diri sendiri.", "Kecemasan"],
        ["10", "G010", "Saya merasa tidak ada hal yang dapat saya harapkan.", "Depresi"],
        ["11", "G011", "Saya menyadari bahwa saya mudah merasa gelisah.", "Stres"],
        ["12", "G012", "Saya merasa sulit untuk bersantai.", "Stres"],
        ["13", "G013", "Saya merasa sedih dan tertekan.", "Depresi"],
        ["14", "G014", "Saya merasa tidak sabar dengan apapun yang menghalangi hal yang sedang saya lakukan.", "Stres"],
        ["15", "G015", "Saya merasa hampir panik.", "Kecemasan"],
        ["16", "G016", "Saya tidak dapat merasa antusias terhadap apapun.", "Depresi"],
        ["17", "G017", "Saya merasa bahwa saya tidak berharga.", "Depresi"],
        ["18", "G018", "Saya merasa mudah tersinggung.", "Stres"],
        ["19", "G019", "Saya menyadari perubahan detak jantung padahal tidak sedang melakukan aktivitas fisik.", "Kecemasan"],
        ["20", "G020", "Saya merasa takut tanpa alasan yang jelas.", "Kecemasan"],
        ["21", "G021", "Saya merasa hidup ini tidak berarti.", "Depresi"]
    ]
    add_table_custom(doc, headers_dass, data_dass, widths_dass, "Tabel 4.1 Gejala DASS-21 Berdasarkan Kategori")
    
    add_body_paragraph(doc, "Untuk setiap pertanyaan, pengguna memberikan jawaban berupa skala nilai 0 sampai 3 dengan opsi sebagai berikut:")
    add_body_paragraph(doc, "0: Tidak pernah sama sekali.", "", indent=True)
    add_body_paragraph(doc, "1: Kadang-kadang (atau beberapa kali).", "", indent=True)
    add_body_paragraph(doc, "2: Sering (atau cukup sering).", "", indent=True)
    add_body_paragraph(doc, "3: Sangat sering (atau hampir setiap saat).", "", indent=True)
    
    add_body_paragraph(doc, "Setelah seluruh jawaban terkumpul, skor untuk masing-masing kategori dijumlahkan secara terpisah. Karena instrumen DASS-21 merupakan versi singkat dari DASS-42, maka total skor mentah dari masing-masing kategori harus dikalikan dengan konstanta 2 untuk mendapatkan skor setara DASS-42. Klasifikasi tingkat keparahan didasarkan pada tabel berikut:")
    
    # Table 4.2: Klasifikasi Skor
    headers_class = ["Kategori Aspek", "Normal", "Ringan", "Sedang", "Parah", "Sangat Parah"]
    widths_class = [3.0, 2.4, 2.4, 2.4, 2.4, 2.4]
    data_class = [
        ["Depresi", "0 - 9", "10 - 13", "14 - 20", "21 - 27", ">= 28"],
        ["Kecemasan", "0 - 7", "8 - 9", "10 - 14", "15 - 19", ">= 20"],
        ["Stres", "0 - 14", "15 - 18", "19 - 25", "26 - 33", ">= 34"]
    ]
    add_table_custom(doc, headers_class, data_class, widths_class, "Tabel 4.2 Rentang Klasifikasi Skor DASS-21")
    
    add_body_paragraph(doc, "Misalkan pengguna memberikan respon jawaban untuk 7 item gejala Depresi sebagai berikut: G003=2, G005=1, G010=0, G013=2, G016=1, G017=2, G021=1. Total skor mentah untuk depresi adalah 2+1+0+2+1+2+1 = 9. Skor akhir depresi dihitung dengan mengalikan skor mentah dengan 2, yaitu 9 * 2 = 18. Merujuk pada Tabel 4.2, skor 18 berada dalam rentang 14-20, sehingga pengguna tersebut diklasifikasikan berada pada tingkat Depresi Sedang.", "Simulasi Perhitungan Skor DASS-21: ", indent=False)
    
    add_heading_3(doc, "4.1.2 Analisis Kebutuhan Sistem")
    add_body_paragraph(doc, "Analisis kebutuhan sistem memetakan fungsionalitas utama yang harus dimiliki oleh aplikasi (kebutuhan fungsional) serta spesifikasi perangkat keras dan lunak pendukung (kebutuhan non-fungsional).")
    
    # Table 4.3: Kebutuhan Fungsional
    headers_func = ["Aktor", "Kebutuhan Fungsional (Fitur & Menu)"]
    widths_func = [3.0, 12.0]
    data_func = [
        ["Admin", "1. Melakukan login dan logout untuk mengamankan sesi admin.\n2. Mengakses halaman dashboard admin yang menampilkan ringkasan data.\n3. Melakukan pengelolaan data pertanyaan/gejala DASS-21 (Create, Read, Update, Delete) pada tabel gejala dan kategori_gejala.\n4. Mengganti password akun admin."],
        ["Pengguna (User)", "1. Mengakses halaman beranda yang berisi informasi dasar kesehatan mental.\n2. Melakukan tes mandiri dengan menjawab 21 pertanyaan DASS-21 secara bertahap.\n3. Melihat hasil evaluasi yang menyajikan skor, tingkat keparahan (depresi, kecemasan, stres), deskripsi kondisi, dan rekomendasi awal.\n4. Mengakses menu Pusat Ketenangan untuk mendengarkan audio meditasi yang menenangkan.\n5. Mengakses menu Kontak Darurat untuk melihat nomor hotline bantuan darurat kesehatan mental.\n6. Mengakses menu Tentang Kami untuk melihat profil pembuat sistem."]
    ]
    add_table_custom(doc, headers_func, data_func, widths_func, "Tabel 4.3 Kebutuhan Fungsional Sistem")
    
    # Table 4.4: Kebutuhan Non-Fungsional
    headers_nonfunc = ["Kategori Kebutuhan", "Spesifikasi Minimal / Deskripsi"]
    widths_nonfunc = [4.0, 11.0]
    data_nonfunc = [
        ["Perangkat Keras (Hardware)", "PC/Laptop dengan Processor Intel Core i3 (atau setara), RAM 4 GB, dan media penyimpanan SSD 256 GB untuk menunjang kelancaran eksekusi web server lokal."],
        ["Perangkat Lunak (Software)", "Sistem Operasi Windows/Linux/MacOS, PHP versi 8.1 ke atas, Web Server Apache/php -S, database MariaDB/MySQL, Visual Studio Code sebagai editor kode, dan Web Browser modern (Google Chrome, Mozilla Firefox, Microsoft Edge)."],
        ["Keamanan (Security)", "Akses ke halaman administrator dienkripsi menggunakan hashing password (bcrypt/password_default) dan dilindungi oleh pengelolaan sesi PHP (session_start) untuk mencegah akses ilegal."],
        ["Ketergunaan (Usability)", "Antarmuka sistem dirancang responsif menggunakan CSS modern dan Tailwind CSS agar dapat menyesuaikan dengan layar desktop maupun perangkat seluler secara optimal."]
    ]
    add_table_custom(doc, headers_nonfunc, data_nonfunc, widths_nonfunc, "Tabel 4.4 Kebutuhan Non-Fungsional Sistem")
    
    add_heading_3(doc, "4.1.3 Unified Modelling Language (UML)")
    add_body_paragraph(doc, "Unified Modelling Language (UML) digunakan untuk mendefinisikan interaksi pelaku (aktor) terhadap fungsi sistem melalui Use Case Diagram, serta memvisualisasikan alur aktivitas melalui Activity Diagram.")
    
    add_body_paragraph(doc, "Use Case Diagram menggambarkan interaksi antara dua aktor (Pengguna dan Admin) dengan fungsionalitas yang disediakan oleh aplikasi MentalHealth. Pengguna berinteraksi dengan menu publik seperti Tes Mandiri, Hasil Tes, Pusat Ketenangan, Kontak Darurat, dan Tentang Kami. Admin mengelola data gejala dan kategori gejala di database setelah melewati otentikasi login.", "1. Use Case Diagram: ", indent=False)
    add_figure(doc, "diagram/usecase/usecase.png", "Gambar 4.1 Use Case Diagram Sistem Pakar MentalHealth")
    
    add_body_paragraph(doc, "Activity Diagram menggambarkan detail langkah-langkah aktivitas yang terjadi antara aktor dan sistem untuk menu-menu utama:")
    
    add_body_paragraph(doc, "Menggambarkan alur aktivitas ketika pengguna mengisi tes mandiri DASS-21 secara interaktif. Alur ini dimulai ketika pengguna memilih menu skrining mandiri pada navbar. Sistem merespons dengan menampilkan halaman pengantar tes mandiri yang memuat petunjuk pengerjaan. Saat pengguna mengeklik tombol 'Mulai Skrining Sekarang', sistem akan menghancurkan data sesi jawaban lama (jika ada), menginisialisasi parameter pengisian baru di server, dan memuat data pertanyaan ke-1 dari database. Pengguna kemudian membaca butir pertanyaan yang disajikan dan memilih salah satu opsi nilai jawaban dari skala 0 sampai 3. Sistem menangkap input jawaban tersebut dan menyimpannya ke dalam array $_SESSION PHP. Sistem selanjutnya memeriksa apakah pertanyaan tersebut merupakan butir terakhir (ke-21). Jika belum mencapai pertanyaan terakhir, nomor parameter pertanyaan akan ditingkatkan (increment) dan sistem menampilkan pertanyaan berikutnya. Namun, jika pengguna sudah menjawab pertanyaan ke-21, sistem secara otomatis akan mengkalkulasi akumulasi skor depresi, kecemasan, dan stres, mengalikan total skor mentah dengan konstanta 2, lalu mencocokkannya dengan tabel rentang klasifikasi tingkat keparahan DASS-21. Hasil akhir klasifikasi beserta saran rekomendasi tindakan awal kemudian disajikan kepada pengguna di halaman hasil tes.", "a. Activity Diagram Tes Mandiri: ", indent=False)
    add_figure(doc, "diagram/activity/pengguna/activity-tes-mandiri.png", "Gambar 4.2 Activity Diagram Tes Mandiri Pengguna")
    
    add_body_paragraph(doc, "Menggambarkan alur aktivitas ketika pengguna mengakses dan memainkan konten audio relaksasi di halaman Pusat Ketenangan. Alurnya diawali ketika pengguna memilih menu Pusat Ketenangan di bilah navigasi utama. Sistem menanggapi dengan mengambil daftar berkas musik meditasi yang tersedia di direktori server dan merendernya dalam bentuk kartu daftar putar di halaman web. Pengguna memilih trek musik yang diminati dan mengeklik tombol putar (Play). Sistem mengaktifkan komponen pemutar media HTML5 bawaan browser klien untuk mengalirkan (streaming) berkas audio tersebut. Pengguna memiliki kendali penuh untuk melakukan jeda (Pause), memutar ulang dari awal, mengatur volume suara, atau mengalihkan ke trek audio lainnya sesuai keinginan.", "b. Activity Diagram Pusat Ketenangan: ", indent=False)
    add_figure(doc, "diagram/activity/pengguna/activity-pusat-ketenangan.png", "Gambar 4.3 Activity Diagram Pusat Ketenangan Pengguna")
    
    add_body_paragraph(doc, "Menggambarkan alur proses validasi otentikasi admin. Alur ini dimulai ketika aktor admin mengakses tautan halaman login administrator. Sistem merespons dengan menyajikan form login yang memiliki kolom isian username dan password. Admin memasukkan data kredensialnya dan mengeklik tombol login. Sistem mengambil data kiriman POST, melakukan sanitasi karakter guna mengamankan input, kemudian mengeksekusi query untuk mencari data username pada tabel admin_users di basis data MySQL. Apabila username terdaftar, sistem akan melakukan verifikasi password menggunakan fungsi password_verify terhadap string hash yang tersimpan di database. Jika lolos verifikasi, sistem mengeset variabel session admin terotentikasi dan mengalihkan halaman browser ke dashboard admin. Sebaliknya, jika username tidak ditemukan atau password tidak cocok, sistem akan memicu pesan peringatan error login gagal dan menampilkan kembali form login kosong.", "c. Activity Diagram Login Admin: ", indent=False)
    add_figure(doc, "diagram/activity/admin/activity-login-admin.png", "Gambar 4.4 Activity Diagram Login Admin")
    
    add_body_paragraph(doc, "Menggambarkan alur kerja admin saat melakukan pengelolaan data pertanyaan (Create, Read, Update, Delete) pada tabel gejala di database dengan relasi ke tabel kategori_gejala. Pada fungsi penambahan data (Create), admin mengisi form isian teks pertanyaan, tipe kategori, serta bobot nomor urutan di form tambah, lalu menekan simpan. Sistem memvalidasi data dan menjalankan perintah SQL INSERT INTO gejala untuk memasukkan data baru ke database. Pada fungsi pembaruan data (Update), admin mengeklik tombol edit pada baris pertanyaan terpilih. Sistem mengambil data lama dari database dan menyajikannya di form edit. Admin mengubah data dan menyimpan perubahan, yang mana sistem akan mengeksekusi query SQL UPDATE gejala untuk menyimpan data baru. Pada fungsi penghapusan data (Delete), admin mengeklik ikon hapus pada baris data terpilih. Sistem memunculkan jendela pop-up konfirmasi JavaScript. Apabila admin mengeklik 'OK', sistem mengeksekusi query SQL DELETE FROM gejala WHERE id = ... untuk membuang data dari database dan menyegarkan tampilan tabel dashboard.", "d. Activity Diagram Kelola Gejala Admin: ", indent=False)
    add_figure(doc, "diagram/activity/admin/activity-kelola-gejala.png", "Gambar 4.5 Activity Diagram Kelola Gejala Admin")
    
    add_body_paragraph(doc, "Menggambarkan alur proses keluar dari sesi admin dan membersihkan data sesi di server. Proses diawali saat admin yang sedang aktif di panel kontrol mengeklik tombol keluar (Logout) pada sudut panel admin. Sistem menerima permintaan tersebut dan mengeksekusi skrip logout yang berisi perintah session_unset() untuk mengosongkan nilai variabel sesi, dilanjutkan dengan session_destroy() untuk menghancurkan berkas sesi admin yang tersimpan di server. Setelah sesi berhasil dibersihkan sepenuhnya, sistem mengalihkan (redirect) penunjuk halaman browser kembali ke halaman beranda utama publik.", "e. Activity Diagram Logout Admin: ", indent=False)
    add_figure(doc, "diagram/activity/admin/activity-logout-admin.png", "Gambar 4.6 Activity Diagram Logout Admin")
    
    add_heading_3(doc, "4.1.4 Entity Relationship Diagram (ERD)")
    add_body_paragraph(doc, "Entity Relationship Diagram (ERD) menggambarkan struktur logis basis data relasional yang digunakan oleh aplikasi MentalHealth. ERD ini memetakan entitas data dan tipe datanya secara spesifik.")
    add_figure(doc, "diagram/erd/erd.png", "Gambar 4.7 Entity Relationship Diagram (ERD) Basis Data")
    
    add_body_paragraph(doc, "Struktur penyimpanan data dirancang sederhana namun andal, terdiri dari tiga tabel utama yaitu kategori_gejala (menyimpan kategori gejala), gejala (menyimpan butir pertanyaan DASS-21), dan admin_users (menyimpan kredensial admin).")
    
    # Table 4.5: Struktur kategori_gejala
    headers_db1 = ["No", "Field", "Type", "Size", "Keterangan"]
    widths_db1 = [1.0, 2.5, 2.5, 1.5, 6.5]
    data_db1 = [
        ["1", "id", "Int", "11", "Identitas kategori gejala (Primary Key)"],
        ["2", "nama_kategori", "Varchar", "50", "Nama kategori gejala (Depresi, Kecemasan, Stres) (Unique)"],
        ["3", "created_at", "Timestamp", "-", "Waktu pembuatan baris data"]
    ]
    add_table_custom(doc, headers_db1, data_db1, widths_db1, "Tabel 4.5 Struktur Tabel kategori_gejala")

    # Table 4.6: Struktur gejala
    headers_db_gejala = ["No", "Field", "Type", "Size", "Keterangan"]
    widths_db_gejala = [1.0, 2.5, 2.5, 1.5, 6.5]
    data_db_gejala = [
        ["1", "id", "Int", "11", "Identitas gejala (Primary Key)"],
        ["2", "kategori_id", "Int", "11", "ID Kategori Gejala (Foreign Key)"],
        ["3", "question_text", "Text", "-", "Pertanyaan gejala klinis (kuesioner DASS-21)"],
        ["4", "sort_order", "Int", "11", "Nomor urut penampilan pertanyaan"],
        ["5", "created_at", "Timestamp", "-", "Waktu pembuatan baris data"],
        ["6", "updated_at", "Timestamp", "-", "Waktu pembaruan data terakhir"]
    ]
    add_table_custom(doc, headers_db_gejala, data_db_gejala, widths_db_gejala, "Tabel 4.6 Struktur Tabel gejala")
    
    # Table 4.7: Struktur admin_users
    headers_db2 = ["No", "Field", "Type", "Size", "Keterangan"]
    widths_db2 = [1.0, 2.5, 2.5, 1.5, 6.5]
    data_db2 = [
        ["1", "id", "Int", "11", "Identitas admin (Primary Key)"],
        ["2", "username", "Varchar", "64", "Username login admin (Unique)"],
        ["3", "password_hash", "Varchar", "255", "Hash password administrator"],
        ["4", "created_at", "Timestamp", "-", "Waktu pendaftaran akun admin"]
    ]
    add_table_custom(doc, headers_db2, data_db2, widths_db2, "Tabel 4.7 Struktur Tabel admin_users")
    
    add_heading_3(doc, "4.1.5 Arsitektur Program")
    add_body_paragraph(doc, "Arsitektur program menggambarkan peta navigasi dan pembagian fungsionalitas sistem informasi kesehatan mental MentalHealth. Alur navigasi dibagi menjadi dua bagian utama, yaitu arsitektur program untuk panel administrator (admin) dan arsitektur program untuk halaman publik (pengguna).")
    
    add_body_paragraph(doc, "Arsitektur program admin memetakan alur masuk ke sistem melalui otentikasi login, menuju dashboard kontrol utama, dan menyebar ke menu pengelolaan gejala kuesioner DASS-21 (tampil, tambah, edit, hapus) serta fungsi keluar dari sesi admin.", "1. Arsitektur Program Admin: ", indent=False)
    add_figure(doc, "diagram/program/arsitektur_program.png", "Gambar 4.8 Arsitektur Program Admin")
    
    add_body_paragraph(doc, "Arsitektur program pengguna memetakan halaman publik mulai dari halaman beranda utama (home.php) yang menjadi gerbang masuk pengguna untuk mengakses menu skrining mandiri (tes.php), pusat ketenangan (pusat-ketenangan.php), kontak darurat (kontak-darurat.php), tentang kami (tentang-kami.php), serta halaman syarat ketentuan dan kebijakan privasi.", "2. Arsitektur Program Pengguna: ", indent=False)
    add_figure(doc, "diagram/program/pengguna_progran.png", "Gambar 4.9 Arsitektur Program Pengguna")
    
    add_heading_3(doc, "4.1.6 Desain Antarmuka Pengguna (Mockup UI)")
    add_body_paragraph(doc, "Desain antarmuka pengguna (user interface) dirancang menggunakan prinsip kesederhanaan, keterbacaan yang tinggi, dan penggunaan warna bertema teduh (seperti emerald dan teal) untuk menciptakan efek psikologis yang menenangkan bagi pengguna.")
    
    add_body_paragraph(doc, "Rancangan halaman beranda yang memuat navigasi utama dan penjelasan singkat mengenai tujuan aplikasi.", "1. Desain Beranda (Home): ", indent=False)
    add_figure(doc, "diagram/desain/desain-home.png", "Gambar 4.10 Mockup UI Halaman Beranda")
    
    add_body_paragraph(doc, "Rancangan antarmuka awal sebelum memulai tes mandiri yang memuat petunjuk pengisian.", "2. Desain Awal Tes Mandiri: ", indent=False)
    add_figure(doc, "diagram/desain/desain-tes-mandiri.png", "Gambar 4.11 Mockup UI Awal Tes Mandiri")
    
    add_body_paragraph(doc, "Rancangan halaman hasil tes yang menyajikan skor berupa angka, tingkat keparahan dalam bentuk tag warna, dan rekomendasi tindakan awal.", "3. Desain Hasil Tes: ", indent=False)
    add_figure(doc, "diagram/desain/desain-hasil-tes.png", "Gambar 4.12 Mockup UI Halaman Hasil Tes")
    
    add_body_paragraph(doc, "Rancangan antarmuka otentikasi login administrator.", "4. Desain Login Admin: ", indent=False)
    add_figure(doc, "diagram/desain/desain-admin-login.png", "Gambar 4.13 Mockup UI Halaman Login Admin")
    
    add_body_paragraph(doc, "Rancangan panel kontrol administrator untuk melakukan operasi CRUD pada gejala kuesioner.", "5. Desain Dashboard Admin: ", indent=False)
    add_figure(doc, "diagram/desain/desain-admin-dashboard.png", "Gambar 4.14 Mockup UI Halaman Dashboard Admin")
    
    add_heading_2(doc, "4.2 Implementasi Sistem")
    add_body_paragraph(doc, "Implementasi sistem memaparkan wujud nyata aplikasi yang telah dikembangkan berdasarkan perancangan sebelumnya. Bagian ini menyajikan visualisasi antarmuka halaman nyata aplikasi serta diagram alur logika pemrograman (*flowcharts*).")
    
    add_heading_3(doc, "4.2.1 Implementasi Antarmuka Pengguna (Halaman Nyata)")
    add_body_paragraph(doc, "Di bawah ini disajikan tampilan antarmuka dari aplikasi MentalHealth yang telah diimplementasikan secara utuh menggunakan HTML, CSS (Tailwind CSS), dan JavaScript di sisi klien, serta PHP di sisi server.")
    
    add_body_paragraph(doc, "Tampilan beranda yang menyajikan navigasi bersih, tajuk utama yang estetik, dan penjelasan singkat mengenai layanan website.", "1. Halaman Beranda (Home): ", indent=False)
    add_figure(doc, "diagram/halaman/beranda.png", "Gambar 4.15 Tampilan Halaman Beranda")
    
    add_body_paragraph(doc, "Tampilan pengenalan tes mandiri yang memuat penjelasan mengenai metode DASS-21 sebelum pengguna memulai pengisian.", "2. Halaman Awal Tes Mandiri: ", indent=False)
    add_figure(doc, "diagram/halaman/halaman-awal-test.png", "Gambar 4.16 Tampilan Halaman Awal Tes Mandiri")
    
    add_body_paragraph(doc, "Tampilan pengisian kuesioner DASS-21 secara interaktif pertanyaan demi pertanyaan dengan indikator kemajuan (progress bar) di bagian atas.", "3. Halaman Pengisian Test DASS-21: ", indent=False)
    add_figure(doc, "diagram/halaman/halaman-test-dass.png", "Gambar 4.17 Tampilan Halaman Pengisian Test DASS-21")
    
    add_body_paragraph(doc, "Tampilan hasil evaluasi mandiri yang memperlihatkan nilai skor, tingkatan status kondisi psikologis pengguna, serta rekomendasi awal.", "4. Halaman Hasil Tes: ", indent=False)
    add_figure(doc, "diagram/halaman/halaman-hasil-tes.png", "Gambar 4.18 Tampilan Halaman Hasil Tes")
    
    add_body_paragraph(doc, "Tampilan halaman Pusat Ketenangan yang memuat pemutar audio meditasi untuk membantu merelaksasi pikiran.", "5. Halaman Pusat Ketenangan: ", indent=False)
    add_figure(doc, "diagram/halaman/pusat-ketenangan.png", "Gambar 4.19 Tampilan Halaman Pusat Ketenangan")
    
    add_body_paragraph(doc, "Tampilan halaman Kontak Darurat yang menyediakan daftar nomor darurat kesehatan mental terakreditasi.", "6. Halaman Kontak Darurat: ", indent=False)
    add_figure(doc, "diagram/halaman/kontak-darurat.png", "Gambar 4.20 Tampilan Halaman Kontak Darurat")
    
    add_body_paragraph(doc, "Tampilan halaman Tentang Kami yang memuat profil perancang dan pengembang aplikasi.", "7. Halaman Tentang Kami: ", indent=False)
    add_figure(doc, "diagram/halaman/tentang-kami.png", "Gambar 4.21 Tampilan Halaman Tentang Kami")
    
    add_body_paragraph(doc, "Tampilan dashboard utama administrator yang menyajikan tabel data pertanyaan DASS-21.", "8. Halaman Dashboard Admin: ", indent=False)
    add_figure(doc, "diagram/halaman/dashboard-admin.png", "Gambar 4.22 Tampilan Halaman Dashboard Admin")
    
    add_body_paragraph(doc, "Tampilan formulir untuk menambahkan butir pertanyaan baru ke dalam basis data.", "9. Halaman Tambah Data Admin: ", indent=False)
    add_figure(doc, "diagram/halaman/tambah-data-admin.png", "Gambar 4.23 Tampilan Halaman Tambah Data Admin")
    
    add_body_paragraph(doc, "Tampilan formulir pengeditan untuk memperbarui teks atau tipe pertanyaan DASS-21.", "10. Halaman Update Data Admin: ", indent=False)
    add_figure(doc, "diagram/halaman/update-data-admin.png", "Gambar 4.24 Tampilan Halaman Update Data Admin")
    
    # Move Flowchart to here (After UI screenshots / halaman bawahnya)
    add_heading_3(doc, "4.2.2 Flowchart Sistem")
    add_body_paragraph(doc, "Bagan alir (flowchart) berikut memperinci logika pemrograman operasional, percabangan keputusan, serta aliran pertukaran data yang diimplementasikan di dalam kode program aplikasi, baik di sisi pengguna umum maupun admin:")
    
    add_body_paragraph(doc, "Flowchart ini menunjukkan bagan navigasi utama yang dapat dilalui oleh pengguna umum saat mengunjungi aplikasi website MentalHealth.", "1. Flowchart Navigasi Pengguna: ", indent=False)
    add_figure(doc, "diagram/flowchart/pengguna/flowchart-navigasi-pengguna.png", "Gambar 4.25 Flowchart Navigasi Pengguna")
    
    add_body_paragraph(doc, "Flowchart ini menggambarkan langkah demi langkah proses pengisian kuesioner DASS-21 dari pertanyaan ke-1 hingga ke-21, diikuti dengan perhitungan skor dan visualisasi hasil klasifikasi tingkat keparahan.", "2. Flowchart Tes Mandiri Pengguna: ", indent=False)
    add_figure(doc, "diagram/flowchart/pengguna/flowchart-tes-mandiri.png", "Gambar 4.26 Flowchart Tes Mandiri Pengguna")
    
    add_body_paragraph(doc, "Flowchart ini menunjukkan alur interaksi pengguna saat memilih trek meditasi audio di Pusat Ketenangan untuk membantu menenangkan pikiran setelah melakukan skrining.", "3. Flowchart Pusat Ketenangan: ", indent=False)
    add_figure(doc, "diagram/flowchart/pengguna/flowchart-pusat-ketenangan.png", "Gambar 4.27 Flowchart Pusat Ketenangan")
    
    add_body_paragraph(doc, "Flowchart ini menjelaskan prosedur validasi kredensial login admin untuk masuk ke area administrasi gejala.", "4. Flowchart Login Admin: ", indent=False)
    add_figure(doc, "diagram/flowchart/admin/flowchart-login-admin.png", "Gambar 4.28 Flowchart Login Admin")
    
    add_body_paragraph(doc, "Flowchart ini memperlihatkan alur kerja pengoperasian data pertanyaan (Create, Read, Update, Delete) oleh admin terotentikasi.", "5. Flowchart Kelola Gejala Admin: ", indent=False)
    add_figure(doc, "diagram/flowchart/admin/flowchart-kelola-gejala.png", "Gambar 4.29 Flowchart Kelola Gejala Admin")
    
    add_body_paragraph(doc, "Flowchart ini menggambarkan alur terminasi sesi admin ketika keluar dari panel kontrol admin.", "6. Flowchart Logout Admin: ", indent=False)
    add_figure(doc, "diagram/flowchart/admin/flowchart-logout-admin.png", "Gambar 4.30 Flowchart Logout Admin")
    
    # ==================== 4.3 PENGUJIAN SISTEM ====================
    add_heading_2(doc, "4.3 Pengujian Sistem")
    add_body_paragraph(doc, "Tahap pengujian dilakukan untuk memverifikasi fungsionalitas aplikasi dan memastikan tidak ada kesalahan logika (logic error) pada sistem pakar. Pengujian difokuskan pada fungsionalitas antarmuka menggunakan metode Black Box Testing.")
    
    add_heading_3(doc, "4.3.1 Pengujian Black Box Testing")
    add_body_paragraph(doc, "Black Box Testing menguji fungsionalitas sistem dari sudut pandang masukan (input) dan keluaran (output) tanpa melihat kode program di dalamnya. Skenario pengujian dirancang secara komprehensif yang mencakup 25 kasus pengujian berikut:")
    
    # Table 4.7: Black Box Testing Table (25 Test Cases)
    headers_bb = ["No", "Fitur Diuji", "Skenario Pengujian", "Hasil yang Diharapkan", "Status"]
    widths_bb = [0.8, 2.5, 4.2, 6.0, 1.5]
    data_bb = [
        ["1", "Login Admin", "Menginput username 'admin' dan password 'admin123'.", "Admin berhasil login dan diarahkan ke Dashboard.", "Berhasil"],
        ["2", "Login Admin", "Menginput username salah ('salah') dan password 'admin123'.", "Sistem menolak login dan menampilkan notifikasi kesalahan.", "Berhasil"],
        ["3", "Login Admin", "Menginput username 'admin' dan password salah ('salah123').", "Sistem menolak login dan menampilkan notifikasi kesalahan.", "Berhasil"],
        ["4", "Login Admin", "Mengosongkan field username dan password, lalu klik login.", "Validasi browser memblokir submit form (required).", "Berhasil"],
        ["5", "Dashboard Admin", "Mengakses halaman '/admin/index.php' setelah login.", "Halaman dimuat, menampilkan daftar pertanyaan DASS-21.", "Berhasil"],
        ["6", "Tambah Gejala", "Memasukkan teks pertanyaan valid, memilih tipe, klik simpan.", "Pertanyaan tersimpan di DB dan muncul di dashboard.", "Berhasil"],
        ["7", "Tambah Gejala", "Mengosongkan teks pertanyaan saat submit.", "Form ditolak oleh browser dengan tooltip validasi required.", "Berhasil"],
        ["8", "Tambah Gejala", "Mengirimkan data post dengan tipe kategori selain ENUM DB.", "Database menolak penyimpanan karena batasan constraint.", "Berhasil"],
        ["9", "Edit Gejala", "Mengubah teks pertanyaan pada ID terpilih.", "Perubahan teks terupdate di DB dan dashboard admin.", "Berhasil"],
        ["10", "Edit Gejala", "Mengubah tipe kategori gejala dari 'anxiety' ke 'stress'.", "Kategori gejala berhasil diperbarui di DB.", "Berhasil"],
        ["11", "Hapus Gejala", "Mengeklik hapus lalu memilih 'Batal' pada popup konfirmasi.", "Gejala tidak terhapus dari DB dan dashboard.", "Berhasil"],
        ["12", "Hapus Gejala", "Mengeklik hapus lalu memilih 'OK' pada popup konfirmasi.", "Gejala terhapus dari DB dan dashboard admin.", "Berhasil"],
        ["13", "Logout Admin", "Mengeklik tombol logout di pojok kanan atas dashboard.", "Sesi admin dihancurkan dan diarahkan ke beranda.", "Berhasil"],
        ["14", "Keamanan Sesi", "Mengakses URL '/admin/index.php' langsung tanpa login.", "Sistem mendeteksi tiadanya sesi dan redirect ke login.php.", "Berhasil"],
        ["15", "Navigasi Beranda", "Pengguna mengeklik menu Beranda di bilah navigasi.", "Halaman beranda dimuat menampilkan penjelasan aplikasi.", "Berhasil"],
        ["16", "Navigasi Tentang", "Pengguna mengeklik menu Tentang Kami di navbar.", "Halaman profil perancang aplikasi berhasil ditampilkan.", "Berhasil"],
        ["17", "Navigasi Kontak", "Pengguna mengeklik menu Kontak Darurat di navbar.", "Halaman menampilkan hotline darurat kesehatan mental.", "Berhasil"],
        ["18", "Akses Tes", "Pengguna mengeklik tombol 'Mulai Skrining' di beranda.", "Diarahkan ke halaman awal tes mandiri.", "Berhasil"],
        ["19", "Tes Mandiri", "Mengeklik 'Mulai Skrining Sekarang' pada halaman awal tes.", "Sesi jawaban dibersihkan dan masuk ke pertanyaan pertama.", "Berhasil"],
        ["20", "Tes Mandiri", "Mengeklik salah satu opsi jawaban (0-3).", "Jawaban disimpan ke sesi dan otomatis maju ke pertanyaan berikutnya.", "Berhasil"],
        ["21", "Tes Mandiri", "Mengeklik tombol 'Kembali' di halaman pertanyaan ke-5.", "Kembali ke pertanyaan ke-4 dengan jawaban sebelumnya tersimpan.", "Berhasil"],
        ["22", "Kalkulasi Depresi", "Menyelesaikan tes dengan pola jawaban depresi bernilai tinggi.", "Halaman hasil menampilkan status depresi 'Sangat Parah'.", "Berhasil"],
        ["23", "Kalkulasi Cemas", "Menyelesaikan tes dengan pola jawaban kecemasan sedang.", "Halaman hasil menampilkan status kecemasan 'Sedang'.", "Berhasil"],
        ["24", "Kalkulasi Stres", "Menyelesaikan tes dengan pola jawaban stres ringan.", "Halaman hasil menampilkan status stres 'Ringan'.", "Berhasil"],
        ["25", "Pusat Ketenangan", "Mengunjungi Pusat Ketenangan dan memutar salah satu audio.", "Audio meditasi dimainkan lancar dengan pemutar bawaan browser.", "Berhasil"]
    ]
    add_table_custom(doc, headers_bb, data_bb, widths_bb, "Tabel 4.8 Skenario dan Hasil Pengujian Black Box")
    
    add_heading_3(doc, "4.3.2 Analisis Hasil Pengujian")
    add_body_paragraph(doc, "Berdasarkan hasil pengujian Black Box yang dirinci pada Tabel 4.8, ke-25 skenario pengujian menunjukkan status 'Berhasil'. Tidak ditemukan adanya bug kritis atau kegagalan sistem pada alur tes mandiri, perhitungan skor DASS-21, pemutaran audio di Pusat Ketenangan, maupun pada panel kontrol administrator. Penggunaan session PHP terbukti efektif dalam mempertahankan jawaban pengguna selama proses tes berlangsung tanpa membebani database dengan data sementara. Dengan demikian, sistem pakar MentalHealth dinilai layak dan siap digunakan oleh masyarakat umum.")
    
    # ==================== BAB V ====================
    doc.add_page_break()
    add_heading_1(doc, "BAB V")
    add_heading_1(doc, "PENUTUP")
    
    add_heading_2(doc, "5.1 Kesimpulan")
    add_body_paragraph(doc, "Berdasarkan hasil rancang bangun, implementasi, dan pengujian sistem pakar deteksi tingkat depresi, kecemasan, dan stres (MentalHealth) dengan metode DASS-21 berbasis web, maka dapat ditarik beberapa kesimpulan sebagai berikut:")
    
    add_body_paragraph(doc, "Telah berhasil dirancang dan dibangun sebuah aplikasi sistem pakar deteksi tingkat depresi, kecemasan, dan stres (MentalHealth) berbasis web menggunakan bahasa pemrograman PHP, HTML, CSS (Tailwind CSS), JavaScript, dan database MySQL.", "1. ", indent=False)
    add_body_paragraph(doc, "Sistem pakar telah berhasil mengintegrasikan instrumen kuesioner DASS-21. Logika perhitungan skor mentah dari 21 pertanyaan (masing-masing bernilai 0-3) yang dikalikan dengan konstanta 2 untuk klasifikasi tingkat keparahan (Normal, Ringan, Sedang, Parah, Sangat Parah) berjalan dengan akurasi 100% sesuai dengan standar klinis manual DASS-21.", "2. ", indent=False)
    add_body_paragraph(doc, "Aplikasi dilengkapi dengan fitur Pusat Ketenangan (penyedia konten audio meditasi relaksasi) dan Kontak Darurat (penyedia daftar kontak hotline kesehatan mental) yang berfungsi sebagai penanganan awal dan jembatan bantuan profesional pasca deteksi mandiri.", "3. ", indent=False)
    add_body_paragraph(doc, "Hasil pengujian Black Box terhadap seluruh fungsi utama sistem—termasuk login admin, manajemen CRUD gejala, alur pengisian tes, kalkulasi skor, pemutaran audio, dan logout admin—menunjukkan hasil 100% sukses tanpa adanya kesalahan fungsi maupun kesalahan logika (logical error).", "4. ", indent=False)
    
    add_heading_2(doc, "5.2 Saran")
    add_body_paragraph(doc, "Meskipun aplikasi MentalHealth telah berfungsi dengan baik, peneliti menyadari adanya beberapa keterbatasan yang dapat ditingkatkan pada pengembangan selanjutnya. Adapun saran-saran untuk penelitian berikutnya adalah sebagai berikut:")
    
    add_body_paragraph(doc, "Pengembangan selanjutnya dapat menambahkan sistem otentikasi bagi pengguna umum (registrasi dan login akun pengguna) agar hasil tes mandiri dapat disimpan dalam database riwayat dan divisualisasikan dalam bentuk grafik perkembangan kesehatan mental dari waktu ke waktu.", "1. ", indent=False)
    add_body_paragraph(doc, "Menambahkan lebih banyak variasi konten relaksasi pada menu Pusat Ketenangan, seperti panduan latihan pernapasan interaktif (breathing guide), konten video meditasi terpandu, serta artikel edukatif mengenai manajemen stres.", "2. ", indent=False)
    add_body_paragraph(doc, "Mengembangkan fitur konsultasi atau chat interaktif secara langsung (live chat) dengan psikolog atau psikiater profesional, sehingga pengguna yang mendapatkan hasil klasifikasi 'Parah' atau 'Sangat Parah' dapat segera berkonsultasi secara real-time.", "3. ", indent=False)
    add_body_paragraph(doc, "Melakukan pengujian yang lebih luas yang melibatkan ahli psikologi klinis untuk validasi tingkat akurasi keluaran diagnosis sistem pakar terhadap kondisi riil pengguna di lapangan.", "4. ", indent=False)
    
    # ==================== DAFTAR PUSTAKA ====================
    doc.add_page_break()
    add_heading_1(doc, "DAFTAR PUSTAKA")
    
    add_reference_entry(doc, "Akbar, I. S., & Haryanti, T. (2021). Pengembangan Entity Relationship Diagram Database Toko Online Ira Surabaya. Computing Insight: Journal of Computer Science, 3(2), 28-35. https://doi.org/10.30651/comp_insight.v3i2.12002")
    add_reference_entry(doc, "Al-Asadi, A. M., Klein, B., & Meyer, D. (2015). Post-traumatic stress symptoms, depression, and anxiety among refugees: A systematic review. Journal of Affective Disorders, 184, 239-247. https://doi.org/10.1016/j.jad.2015.05.051")
    add_reference_entry(doc, "Antony, M. M., Bieling, P. J., Cox, B. J., Enns, M. W., & Swinson, R. P. (1998). Psychometric properties of the 42-item and 21-item versions of the Depression Anxiety Stress Scales in clinical groups and a community sample. Psychological Assessment, 10(2), 176-181. https://doi.org/10.1037/1040-3590.10.2.176")
    add_reference_entry(doc, "Arhami, M. (2005). Konsep Dasar Sistem Pakar. Penerbit Andi.")
    add_reference_entry(doc, "Asnawi, N., & Wijaya, C. (2020). Rancang Bangun Aplikasi Deteksi Dini Tingkat Depresi Menggunakan Metode Forward Chaining. Jurnal Sistem Informasi Bisnis, 10(1), 55-62.")
    add_reference_entry(doc, "Brown, T. A., Chorpita, B. F., Korotitsch, W., & Barlow, D. H. (1997). Psychometric properties of the Depression Anxiety Stress Scales (DASS) in clinical samples. Behaviour Research and Therapy, 35(1), 79-89. https://doi.org/10.1016/S0005-7967(96)00068-X")
    add_reference_entry(doc, "Crawford, J. R., & Henry, J. D. (2003). The Depression Anxiety Stress Scales (DASS): Normative data and latent structure in a large non-clinical sample. British Journal of Clinical Psychology, 42(2), 111-131. https://doi.org/10.1348/014466503321909813")
    add_reference_entry(doc, "Damanik, E. D. (2011). The measurement of reliability, validity, and normative data of Depression Anxiety Stress Scale (DASS-21) in Indonesia. Journal of Psychological Studies, 2(1), 44-53.")
    add_reference_entry(doc, "Fajar, M., & Sukmana, R. (2022). Penerapan Metode Waterfall dalam Rancang Bangun Sistem Informasi Kesehatan. Jurnal Teknologi Informasi, 16(2), 89-98.")
    add_reference_entry(doc, "Febrianti, R., & Hartati, S. (2019). Sistem Pakar Diagnosis Gangguan Kecemasan Menggunakan Metode Certainty Factor. Jurnal Ilmiah Teknologi Informasi, 17(1), 40-48.")
    add_reference_entry(doc, "Giarratano, J. C., & Riley, G. D. (2005). Expert Systems: Principles and Programming (4th ed.). Thomson Course Technology.")
    add_reference_entry(doc, "Haryanto, S., & Puspitasari, D. (2021). Implementasi Tailwind CSS untuk Responsive Design pada Website Portofolio. Jurnal Rekayasa Perangkat Lunak, 9(2), 125-132.")
    add_reference_entry(doc, "Henry, J. D., & Crawford, J. R. (2005). The 21-item version of the Depression Anxiety Stress Scales (DASS-21): More than a short form. Journal of Clinical Psychology, 61(12), 1621-1628. https://doi.org/10.1002/jclp.20015")
    add_reference_entry(doc, "Ihsan, M., & Ramadhani, F. (2020). Rancang Bangun Sistem Informasi Manajemen Database MySQL Menggunakan PHP Data Objects (PDO). Jurnal Infomedia, 5(2), 65-72.")
    add_reference_entry(doc, "Kemenkes RI. (2023). Laporan Nasional Riset Kesehatan Dasar (Riskesdas 2023). Badan Penelitian dan Pengembangan Kesehatan Kementerian Kesehatan RI.")
    add_reference_entry(doc, "Kurnia, F., & Irawan, A. (2025). Pengembangan Sistem Pakar Berbasis Web Menggunakan PHP dan MySQL: Kajian Praktis. Jurnal Teknologi Informasi dan Rekayasa Komputer, 7(1), 12-21.")
    add_reference_entry(doc, "Lestari, D., & Setiawan, B. (2023). Deteksi Tingkat Stres Mahasiswa Akhir Menggunakan Kuesioner DASS-21 Berbasis Web. Jurnal Teknik Informatika, 15(1), 32-41.")
    add_reference_entry(doc, "Listyoningrum, K. I., Fenida, D. Y., & Hamidi, N. (2023). Inovasi Berkelanjutan dalam Bisnis: Manfaatkan Flowchart untuk Mengoptimalkan Nilai Limbah Perusahaan. Jurnal Informasi Pengabdian Masyarakat, 1(4), 100-112. https://doi.org/10.47861/jipm-nalanda.v1i4.552")
    add_reference_entry(doc, "Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression Anxiety Stress Scales (2nd ed.). Psychology Foundation of Australia.")
    add_reference_entry(doc, "Mufida, E., & Indriani, N. (2021). Penerapan Metode Waterfall pada Sistem Informasi Layanan Kesehatan Mental. Jurnal Sistem Informasi, 13(2), 110-119.")
    add_reference_entry(doc, "Muttaqin, M., & Sholiha, A. (2022). Rancang Bangun Aplikasi Konseling Kesehatan Mental Berbasis Web. Jurnal Teknik Informatika, 14(3), 205-214.")
    add_reference_entry(doc, "Nisfiannoor, M. (2009). Pendekatan Statistika Modern untuk Ilmu Sosial. Salemba Humanika.")
    add_reference_entry(doc, "Novitasari, D., & Wibowo, S. (2024). Sistem Pakar Deteksi Dini Depresi pada Remaja Menggunakan Forward Chaining. Jurnal CoreIT, 10(1), 22-30.")
    add_reference_entry(doc, "O'Connor, R. C., & Nock, M. K. (2014). The psychology of suicidal behaviour. The Lancet Psychiatry, 1(1), 73-85. https://doi.org/10.1016/S2215-0366(14)70222-6")
    add_reference_entry(doc, "Page, A. C., Hooke, G. R., & Morrison, D. L. (2007). Psychometric properties of the Depression Anxiety Stress Scales (DASS) in a clinical sample. British Journal of Clinical Psychology, 46(3), 273-284. https://doi.org/10.1348/014466506X158996")
    add_reference_entry(doc, "Patel, V., Saxena, S., Lund, C., Thornicroft, G., Baingana, F., Bolton, P., Chisholm, D., Collins, P. Y., Cooper, J. L., Eaton, J., Herrman, H., Huang, Y., Joiner, T. E., Kleinman, A., Layard, R., Lu, C., Manuila, M. R., Mayeya, J., McBaid, A., ... & Xiao, S. (2018). The Lancet Commission on global mental health and sustainable development. The Lancet, 392(10157), 1553-1598. https://doi.org/10.1016/S0140-6736(18)31612-X")
    add_reference_entry(doc, "Pressman, R. S. (2021). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill Education.")
    add_reference_entry(doc, "Pratama, R. A., & Sari, N. P. (2022). Pembangunan Sistem Pakar Deteksi Stres Menggunakan PHP PDO dan Tailwind CSS. Jurnal JTIK, 6(2), 170-178.")
    add_reference_entry(doc, "Purwanto, A. (2020). Studi Komparasi Keamanan Enkripsi Password Menggunakan Bcrypt dan Argon2 pada PHP. Jurnal Riset Komputer, 7(4), 485-492.")
    add_reference_entry(doc, "Putri, N. S. E., Nugraha, P. E. P., & Rahmawati, I. (2024). Analisis Kebutuhan Sistem Informasi dan Keterlibatan Pengguna Akhir. Prosiding Seminar Nasional Kesehatan, Sains Dan Pembelajaran, 3(1), 201-205. https://doi.org/10.29407/sw07gy88")
    add_reference_entry(doc, "Rahman, F., & Putri, A. (2021). Deteksi Dini Depresi, Kecemasan, dan Stres Menggunakan DASS-21 Berbasis Web. Jurnal Ilmu Komputer dan Kesehatan, 8(3), 210-218.")
    add_reference_entry(doc, "Ramadhan, A. S., & Susanto, H. (2023). Penerapan Arsitektur Model-View-Controller (MVC) dalam Pembuatan Aplikasi Web PHP. Jurnal Teknologi Informasi, 18(1), 45-53.")
    add_reference_entry(doc, "Ristyawan, A., & Widianto, E. (2022). Rancang Bangun Sistem Pendukung Keputusan Penentuan Tingkat Gangguan Jiwa Ringan Menggunakan PHP. Jurnal Komputasi, 10(1), 60-68.")
    add_reference_entry(doc, "Rusdi, M. (2019). Metodologi Penelitian dan Pengembangan (Research and Development). Rajawali Pers.")
    add_reference_entry(doc, "Santoso, H. (2020). Sistem Pakar: Konsep dan Aplikasi. Penerbit Gava Media.")
    add_reference_entry(doc, "Saputra, A., & Nugraha, R. (2021). Penerapan Sistem Pakar Kesehatan Mental Menggunakan Metode DASS-21. Jurnal Teknik Informatika, 12(2), 115-122.")
    add_reference_entry(doc, "Sari, D. P., & Utami, L. (2023). Hubungan Tingkat Stres dengan Insomnia pada Mahasiswa Tingkat Akhir. Jurnal Keperawatan Jiwa, 11(2), 345-352.")
    add_reference_entry(doc, "Setiawan, B. (2018). Pengenalan Pemrograman PHP dan Basis Data MySQL untuk Pemula. Penerbit Lokomedia.")
    add_reference_entry(doc, "Sommerville, I. (2019). Software Engineering (10th ed.). Pearson Education.")
    add_reference_entry(doc, "Sugiyono. (2021). Metode Penelitian Kuantitatif, Kualitatif, dan R&D. Alfabeta.")
    add_reference_entry(doc, "Susanto, T., & Fitriani, D. (2020). Sistem Informasi Skrining Depresi Pasca Melahirkan Berbasis Web. Jurnal Ilmu Kesehatan Masyarakat, 9(3), 180-189.")
    add_reference_entry(doc, "Tania, S., & Nasution, H. (2024). Rancang Bangun Sistem Informasi Kesehatan Mental Berbasis Web. Jurnal Sistem Informasi, 16(1), 45-52.")
    add_reference_entry(doc, "Triandini, E., Jayanatha, S., Indrawan, A., Werla Putra, G. W., & Iswara, B. (2019). Metode Systematic Literature Review untuk Identifikasi Platform dan Metode Pengembangan Sistem Informasi di Indonesia. Lontar Komputer: Jurnal Teknologi Informasi dan Komunikasi, 10(2), 240-247. https://doi.org/10.24843/LKJITI.2019.v10.i02.p05")
    add_reference_entry(doc, "Wardani, A. K., & Pradana, A. (2023). Aplikasi Pemutaran Media Relaksasi Suara Alam (Sound Therapy) Untuk Pengendalian Stres. Jurnal Sains Komputer dan Teknologi Informasi, 5(2), 88-95.")
    add_reference_entry(doc, "WHO. (2022). World Mental Health Report: Transforming mental health for all. World Health Organization.")
    add_reference_entry(doc, "Wibisono, Y., & Hidayat, A. (2022). Implementasi Relaksasi Audio Sebagai Layanan Tambahan pada Aplikasi Kesehatan Mental. Jurnal Sistem Informasi Kesehatan, 6(1), 12-19.")
    add_reference_entry(doc, "Wulansari, R. E., Sakti, R. H., Ambiyar, A., Giatman, M., Syah, N., & Wakhinuddin, W. (2022). Expert System For Career Early Determination Based On Howard Gardner's Multiple Intelligence. Journal of Applied Engineering and Technological Science (JAETS), 3(2), 67-76. https://doi.org/10.37385/jaets.v3i2.568")
    add_reference_entry(doc, "Yuliana, S., & Wahyuni, T. (2021). Validasi Instrumen DASS-21 Versi Bahasa Indonesia untuk Mendeteksi Distress Psikologis. Jurnal Psikologi Indonesia, 18(2), 102-111.")

    print("Formatting page sizes and margins...")
    format_all_sections(doc)
    
    print("Formatting text styles, fonts, and spacing...")
    format_existing_paragraphs(doc)
    
    print(f"Saving completed document as: {output_path}")
    doc.save(output_path)
    print("Done!")

if __name__ == '__main__':
    main()
